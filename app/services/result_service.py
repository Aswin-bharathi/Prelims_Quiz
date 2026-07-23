import io
import pandas as pd
from docx import Document
from app.database import get_db
from app.utils.formatting import format_duration


class ResultService:

    @staticmethod
    def _announcement_key(prefix, quiz_type_id=None):
        if quiz_type_id is not None:
            return f'{prefix}:{quiz_type_id}'
        return prefix

    @staticmethod
    def set_announcement_state(is_released, top_n=None, quiz_type_id=None):
        with get_db() as (_, c):
            released_key = ResultService._announcement_key('results_announced', quiz_type_id)
            c.execute(
                '''INSERT INTO settings (setting_key, setting_value) VALUES (%s, %s)
                   ON DUPLICATE KEY UPDATE setting_value = VALUES(setting_value)''',
                (released_key, '1' if is_released else '0'),
            )
            if top_n is not None:
                top_n_key = ResultService._announcement_key('results_top_n', quiz_type_id)
                c.execute(
                    '''INSERT INTO settings (setting_key, setting_value) VALUES (%s, %s)
                       ON DUPLICATE KEY UPDATE setting_value = VALUES(setting_value)''',
                    (top_n_key, str(top_n)),
                )

    @staticmethod
    def get_announcement_state(quiz_type_id=None):
        with get_db() as (_, c):
            key = ResultService._announcement_key('results_announced', quiz_type_id)
            c.execute('SELECT setting_value FROM settings WHERE setting_key = %s', (key,))
            row = c.fetchone()
            return bool(row and row['setting_value'] == '1')

    @staticmethod
    def get_announcement_top_n(quiz_type_id=None):
        with get_db() as (_, c):
            key = ResultService._announcement_key('results_top_n', quiz_type_id)
            c.execute('SELECT setting_value FROM settings WHERE setting_key = %s', (key,))
            row = c.fetchone()
            if row and row['setting_value'] is not None:
                try:
                    return int(row['setting_value'])
                except (TypeError, ValueError):
                    pass
            return 5

    @staticmethod
    def is_team_selected(lotname, quiz_type_id=None, top_n=None):
        if not lotname:
            return False
        resolved_top_n = top_n if top_n is not None else ResultService.get_announcement_top_n(quiz_type_id)
        leaderboard = ResultService.get_leaderboard(quiz_type_id, resolved_top_n)
        return any(team['lotname'] == lotname for team in leaderboard)

    @staticmethod
    def get_ranked_results(page=1, per_page=10, search='', quiz_type_id=None, top_n=None):
        offset = (page - 1) * per_page

        with get_db() as (_, c):

            where_clause = "WHERE 1=1"
            params = []

            if quiz_type_id is not None:
                where_clause += " AND r.quiz_type_id = %s"
                params.append(quiz_type_id)

            if search:
                where_clause += " AND s.lotname LIKE %s"
                params.append(f"%{search}%")

            # NOTE: Ranking is computed in Python, not via MySQL session
            # variables (@rank := ...). That pattern reads and writes user
            # variables within the same SELECT list, and MySQL does not
            # guarantee left-to-right evaluation order for that - it
            # "usually" works but can silently break depending on the query
            # plan, join order, or indexes used. Sorting is reliable in SQL;
            # ranking/filtering/pagination is done here instead, where
            # order is guaranteed.
            base_query = f"""
                SELECT
                    r.id,
                    r.student_id,
                    r.quiz_type_id,
                    r.score,
                    r.duration,
                    r.total_questions,
                    r.completed_at,
                    qt.name AS quiz_type_name,
                    s.lotname
                FROM results r
                JOIN students s
                    ON s.id = r.student_id
                JOIN quiz_types qt
                    ON qt.id = r.quiz_type_id
                {where_clause}
                ORDER BY
                    r.quiz_type_id,
                    r.score DESC,
                    r.duration ASC,
                    s.lotname ASC
            """

            c.execute(base_query, params)
            all_rows = c.fetchall()

            # Assign rank per quiz_type_id, resetting on quiz_type change,
            # same rank for ties on (score, duration).
            rank = 0
            prev_key = None
            for row in all_rows:
                key = (row['quiz_type_id'], row['score'], row['duration'])
                if key != prev_key:
                    rank += 1
                row['rank'] = rank
                prev_key = key
                row['duration_formatted'] = format_duration(row['duration'])

            if top_n is not None:
                all_rows = [row for row in all_rows if row['rank'] <= top_n]

            # Final display order: by quiz type name, then rank, then team name
            all_rows.sort(key=lambda r: (r['quiz_type_name'], r['rank'], r['lotname']))

            total = len(all_rows)
            results = all_rows[offset:offset + per_page]

            total_pages = max(1, (total + per_page - 1) // per_page)

            return results, total, total_pages

    @staticmethod
    def get_leaderboard(quiz_type_id=None, limit=5):
        with get_db() as (_, c):
            # Ranking logic: prioritize score (DESC), then duration (ASC).
            #
            # NOTE: We deliberately do NOT use MySQL session variables
            # (@rank := ...) to compute the rank in SQL. That pattern reads
            # and writes user variables within the same SELECT list, and
            # MySQL does not guarantee left-to-right evaluation order for
            # that - it "usually" works but can silently break depending on
            # the query plan, join order, or indexes used, which is exactly
            # what was causing the leaderboard cards to show stale/scrambled
            # ranks even after adding CAST(... AS UNSIGNED).
            #
            # Instead: let SQL do what it's reliable at (sorting), and
            # compute the rank in Python, where iteration order is
            # guaranteed and there's no ambiguity.
            query = '''
                SELECT
                    s.lotname,
                    r.score,
                    r.duration,
                    qt.name AS quiz_type_name,
                    r.quiz_type_id,
                    r.total_questions
                FROM results r
                JOIN quiz_types qt ON qt.id = r.quiz_type_id
                JOIN students s ON s.id = r.student_id
            '''
            params = []
            if quiz_type_id is not None:
                query += ' WHERE r.quiz_type_id = %s'
                params.append(quiz_type_id)

            query += ' ORDER BY r.quiz_type_id, r.score DESC, r.duration ASC, s.lotname ASC'

            c.execute(query, params)
            all_rows = c.fetchall()

            # Compute rank in Python: same rank for ties (same quiz_type,
            # score, duration), increments otherwise, resets per quiz_type.
            rank = 0
            prev_key = None
            for row in all_rows:
                key = (row['quiz_type_id'], row['score'], row['duration'])
                if key != prev_key:
                    rank += 1
                row['rank'] = rank
                prev_key = key
                row['duration_formatted'] = format_duration(row['duration'])

            top_rows = [row for row in all_rows if row['rank'] <= limit]
            top_rows.sort(key=lambda r: (r['rank'], r['lotname']))
            return top_rows

    @staticmethod
    def export_word(quiz_type_id=None, top_n=5):
        top = ResultService.get_leaderboard(quiz_type_id, top_n)
        doc = Document()

        title = f'Top {top_n} Teams - All Quizzes'
        if quiz_type_id:
            if top:
                title = f'Top {top_n} Teams - {top[0]["quiz_type_name"]}'
            else:
                from app.services.quiz_type_service import QuizTypeService
                quiz_type = QuizTypeService.get_by_id(quiz_type_id)
                title = f'Top {top_n} Teams - {quiz_type["name"] if quiz_type else "Selected Quiz"}'

        doc.add_heading(title, 0)
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        headers = ['Rank', 'Team Name', 'Quiz Type', 'Score', 'Time Taken']
        for i, h in enumerate(headers):
            table.rows[0].cells[i].text = h
        for row in top:
            cells = table.add_row().cells
            cells[0].text = str(row['rank'])
            cells[1].text = row['lotname']
            cells[2].text = row['quiz_type_name']
            cells[3].text = str(row['score'])
            cells[4].text = row['duration_formatted']
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        return output

    @staticmethod
    def export_excel(quiz_type_id=None, search='', top_n=None):
        # The `top_n` filter should only apply when a specific quiz is selected,
        # matching the behavior of the main results table on the UI.
        effective_top_n = top_n if quiz_type_id is not None else None

        results, _, _ = ResultService.get_ranked_results(
            page=1,
            per_page=100000,
            search=search,
            quiz_type_id=quiz_type_id,
            top_n=effective_top_n
        )
        df = pd.DataFrame([
            {
                'Rank': r['rank'],
                'Team': r['lotname'],
                'Quiz Type': r['quiz_type_name'],
                'Score': r['score'],
                'Total Questions': r['total_questions'],
                'Duration': r['duration_formatted'],
            }
            for r in results
        ])
        output = io.BytesIO()
        with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
            df.to_excel(writer, index=False, sheet_name='Results')
        output.seek(0)
        return output

    @staticmethod
    def unset_announcement_state(quiz_type_id=None):
        """Cancel/reset a release without touching top_n."""
        with get_db() as (_, c):
            released_key = ResultService._announcement_key('results_announced', quiz_type_id)
            c.execute(
                '''INSERT INTO settings (setting_key, setting_value) VALUES (%s, %s)
                   ON DUPLICATE KEY UPDATE setting_value = VALUES(setting_value)''',
                (released_key, '0'),
            )

    @staticmethod
    def get_all_announcement_states():
        """Returns {quiz_type_id: {'released': bool, 'top_n': int}} for every quiz type
        that has ever had a release toggled."""
        with get_db() as (_, c):
            c.execute(
                "SELECT setting_key, setting_value FROM settings "
                "WHERE setting_key LIKE 'results_announced:%'"
            )
            rows = c.fetchall()
            states = {}
            for row in rows:
                try:
                    qid = int(row['setting_key'].split(':', 1)[1])
                except (IndexError, ValueError):
                    continue
                states[qid] = {
                    'released': row['setting_value'] == '1',
                    'top_n': ResultService.get_announcement_top_n(qid),
                }
            return states